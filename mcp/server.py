"""
Trello Deep — MCP-сервер.

Закрывает четыре пробела штатного Trello-коннектора: вложения (список +
скачивание), комментарии, история изменений. Плюс два пишущих инструмента
(комментарий, загрузка файла) с обязательным подтверждением.

Инструменты:
  - trello_fetch_attachment      скачать вложение на диск (read)
  - trello_read_attachment_text  извлечь текст брифа PDF/docx (read)
  - trello_list_attachments      список вложений          (read)
  - trello_get_comments          комментарии карточки      (read)
  - trello_get_card_history      история изменений         (read)
  - trello_add_comment        добавить комментарий      (write, confirm)
  - trello_upload_attachment  прикрепить файл           (write, confirm)

Ключ и токен читаются ИСКЛЮЧИТЕЛЬНО из переменных окружения
TRELLO_KEY и TRELLO_TOKEN (или из .env рядом). Хардкод недопустим.
"""

import asyncio
import json
import mimetypes
import os
import re
from pathlib import Path
from urllib.parse import quote

import httpx
from dotenv import load_dotenv
from mcp.server.fastmcp import FastMCP

# .env авторитетнее переменных окружения процесса: install.ps1/.sh пишут токен
# именно в .env, и он должен побеждать возможный УСТАРЕВШИЙ TRELLO_TOKEN, который
# мог остаться в окружении (иначе плагин молча ходит со старым токеном -> 401).
# Если .env нет — используются системные переменные окружения как есть.
_HERE = Path(__file__).resolve().parent
load_dotenv(_HERE.parent / ".env", override=True)   # .env в корне репозитория
load_dotenv(_HERE / ".env", override=True)          # .env рядом с server.py, если есть

API = "https://api.trello.com/1"
KEY = os.environ.get("TRELLO_KEY")
TOKEN = os.environ.get("TRELLO_TOKEN")

mcp = FastMCP("trello-deep")


# --------------------------------------------------------------------------
# Вспомогательное
# --------------------------------------------------------------------------

class TrelloError(Exception):
    """Ошибка, текст которой безопасно показать агенту/пользователю."""


def _require_creds() -> None:
    if not KEY or not TOKEN:
        raise TrelloError(
            "Не заданы TRELLO_KEY и/или TRELLO_TOKEN. "
            "Скопируй .env.example в .env и подставь ключ и токен."
        )


def _auth_header() -> dict[str, str]:
    # Trello отключил авторизацию через query-параметры для скачивания
    # вложений в январе 2021. Работает только этот заголовок — используем
    # его единообразно на всех эндпоинтах.
    return {
        "Authorization": f'OAuth oauth_consumer_key="{KEY}", oauth_token="{TOKEN}"'
    }


def _normalize_card(card: str) -> str:
    """URL карточки -> shortLink. shortLink и полный id пропускает как есть."""
    m = re.search(r"/c/([A-Za-z0-9]+)", card)
    if m:
        return m.group(1)
    return card.strip()


def _normalize_board(board: str) -> str:
    """URL доски -> shortLink. shortLink и полный id пропускает как есть."""
    m = re.search(r"/b/([A-Za-z0-9]+)", board)
    if m:
        return m.group(1)
    return board.strip()


async def _api_get(path: str, **params):
    """Разовый GET JSON: создаёт клиент, вызывает _get_json. Для читающих тулов."""
    _require_creds()
    async with httpx.AsyncClient(timeout=60) as client:
        return await _get_json(client, path, **params)


async def _api_post(path: str, *, data: dict | None = None, files: dict | None = None):
    """Разовый POST к Trello (пишущие тулы). Авторизация через key/token
    query-параметрами — для JSON-эндпоинтов это допустимо и надёжно для POST."""
    _require_creds()
    params = {"key": KEY, "token": TOKEN}
    async with httpx.AsyncClient(timeout=httpx.Timeout(300.0, connect=30.0)) as client:
        r = await client.post(f"{API}{path}", params=params, data=data, files=files)
        _raise_for_trello(r, path)
        return r.json()


async def _get_json(client: httpx.AsyncClient, path: str, **params):
    """GET JSON с ретраями на 429 (экспоненциальная задержка, до 3 попыток)."""
    delay = 1.0
    for attempt in range(3):
        r = await client.get(f"{API}{path}", params=params, headers=_auth_header())
        if r.status_code == 429 and attempt < 2:
            await asyncio.sleep(delay)
            delay *= 2
            continue
        _raise_for_trello(r, path)
        return r.json()
    _raise_for_trello(r, path)  # последний ответ, если все попытки — 429


def _raise_for_trello(r: httpx.Response, what: str) -> None:
    if r.status_code == 401:
        raise TrelloError(
            "401 Unauthorized. Проверь TRELLO_KEY/TRELLO_TOKEN и что токен "
            "выпущен со scope, включающим read. Скачивание работает только "
            "с OAuth-заголовком (не с query-параметрами)."
        )
    if r.status_code == 404:
        raise TrelloError(f"404 Not Found: {what}. Проверь id карточки/вложения.")
    if r.status_code == 429:
        raise TrelloError("429 Rate limit: превышена частота запросов к Trello.")
    r.raise_for_status()


# --------------------------------------------------------------------------
# Этап 1: скачивание вложения
# --------------------------------------------------------------------------

@mcp.tool()
async def trello_fetch_attachment(
    card: str,
    attachment_id: str,
    dest_dir: str = "./trello_files",
) -> dict:
    """Скачивает вложение карточки Trello в файловую систему и возвращает ПУТЬ к файлу.

    Возвращает путь, а не содержимое: вложения бывают по 20+ МБ и забили бы
    контекст. Прочитай файл по возвращённому пути своими штатными средствами
    (в т.ч. картинки — их видно нативно).

    Аргументы:
        card: URL карточки, shortLink (напр. "KdZqQd8E") или полный id.
        attachment_id: id вложения (см. trello_list_attachments).
        dest_dir: куда сохранить. По умолчанию "./trello_files".

    Возвращает при файле:  {"path", "bytes", "mime", "is_link": false}
    Возвращает при ссылке: {"is_link": true, "url", "name", "path": null}
        (isUpload=false — это внешняя ссылка на Google Drive/YouTube и т.п.,
         скачивать нечего).
    """
    _require_creds()
    card = _normalize_card(card)

    async with httpx.AsyncClient(timeout=httpx.Timeout(300.0, connect=30.0)) as client:
        meta = await _get_json(
            client, f"/cards/{card}/attachments/{attachment_id}", fields="all"
        )

        # Внешняя ссылка — качать нечего.
        if not meta.get("isUpload", False):
            return {
                "is_link": True,
                "url": meta.get("url"),
                "name": meta.get("name"),
                "path": None,
            }

        path, cached = await _download_attachment(client, card, meta, dest_dir)

    mime = meta.get("mimeType") or mimetypes.guess_type(str(path))[0] or "application/octet-stream"
    return {
        "path": str(path),
        "bytes": path.stat().st_size,
        "mime": mime,
        "is_link": False,
        "cached": cached,
    }


async def _stream_download(client: httpx.AsyncClient, url: str, path: Path) -> int:
    """Потоковое скачивание с ретраями на 429. Возвращает размер в байтах."""
    delay = 1.0
    for attempt in range(3):
        try:
            written = 0
            async with client.stream("GET", url, headers=_auth_header()) as r:
                if r.status_code == 429 and attempt < 2:
                    await r.aclose()
                    await asyncio.sleep(delay)
                    delay *= 2
                    continue
                _raise_for_trello(r, str(path))
                with open(path, "wb") as f:
                    async for chunk in r.aiter_bytes(1 << 16):
                        f.write(chunk)
                        written += len(chunk)
            return written
        except httpx.TransportError as e:
            if attempt < 2:
                await asyncio.sleep(delay)
                delay *= 2
                continue
            raise TrelloError(f"Сеть: не удалось скачать {path.name}: {e}") from e
    raise TrelloError(f"Не удалось скачать {path.name} после 3 попыток (429).")


# --------------------------------------------------------------------------
# Кэш вложений: не перекачивать файл, если он не менялся
# --------------------------------------------------------------------------

def _cache_path(dest: Path) -> Path:
    return dest / ".trello_cache.json"


def _read_cache(dest: Path) -> dict:
    p = _cache_path(dest)
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def _write_cache(dest: Path, cache: dict) -> None:
    try:
        _cache_path(dest).write_text(json.dumps(cache, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


_cache_lock = asyncio.Lock()  # защита манифеста кэша при параллельных загрузках


async def _download_attachment(client: httpx.AsyncClient, card: str, meta: dict, dest_dir: str):
    """Скачивает вложение с кэшем. Если файл уже скачан и не менялся
    (совпадают id + размер + дата в манифесте .trello_cache.json), возвращает его
    без повторной загрузки. Возвращает (path, cached: bool)."""
    aid = meta["id"]
    file_name = meta.get("fileName") or meta.get("name") or aid
    size = meta.get("bytes")
    date = meta.get("date")
    dest = Path(dest_dir)
    dest.mkdir(parents=True, exist_ok=True)
    path = dest / file_name

    cache = _read_cache(dest)
    ent = cache.get(aid)
    if (path.exists() and ent and ent.get("bytes") == size and ent.get("date") == date
            and (size is None or path.stat().st_size == size)):
        return path, True  # кэш-хит: файл на месте и не менялся

    url = f"{API}/cards/{card}/attachments/{aid}/download/{quote(file_name)}"
    await _stream_download(client, url, path)
    async with _cache_lock:  # перечитываем и мержим, чтобы параллель не затёрла кэш
        fresh = _read_cache(dest)
        fresh[aid] = {"name": file_name, "bytes": size, "date": date}
        _write_cache(dest, fresh)
    return path, False


# --------------------------------------------------------------------------
# Этап 2: читающие инструменты (list / comments / history)
# --------------------------------------------------------------------------

@mcp.tool()
async def trello_list_attachments(card: str) -> list[dict]:
    """Возвращает список вложений карточки Trello.

    Файлы и внешние ссылки разделяются по признаку `is_upload`:
    is_upload=true — загруженный файл (можно скачать trello_fetch_attachment),
    is_upload=false — внешняя ссылка (Google Drive/YouTube и т.п.).

    Аргументы:
        card: URL карточки, shortLink (напр. "KdZqQd8E") или полный id.

    Возвращает массив объектов: {id, name, mime, bytes, is_upload, url, date}.
    """
    card = _normalize_card(card)
    atts = await _api_get(f"/cards/{card}/attachments", fields="all")
    out = []
    for a in atts:
        is_upload = bool(a.get("isUpload"))
        out.append(
            {
                "id": a.get("id"),
                "name": a.get("fileName") if is_upload else a.get("name"),
                "mime": a.get("mimeType"),
                "bytes": a.get("bytes"),
                "is_upload": is_upload,
                "url": a.get("url"),
                "date": a.get("date"),
            }
        )
    return out


@mcp.tool()
async def trello_get_comments(card: str, limit: int = 1000) -> list[dict]:
    """Возвращает комментарии карточки Trello (то, чего не отдаёт штатный коннектор).

    Аргументы:
        card: URL карточки, shortLink или полный id.
        limit: сколько комментариев вернуть. По умолчанию 1000.
            (Дефолт Trello — 50, поэтому для полной выборки указываем явно.)

    Возвращает массив объектов: {author, date, text}, от новых к старым.
    """
    card = _normalize_card(card)
    actions = await _api_get(
        f"/cards/{card}/actions", filter="commentCard", limit=limit
    )
    return [
        {
            "author": a.get("memberCreator", {}).get("fullName", "—"),
            "date": a.get("date"),
            "text": a.get("data", {}).get("text", ""),
        }
        for a in actions
    ]


@mcp.tool()
async def trello_get_card_history(
    card: str,
    limit: int = 1000,
    types: str | None = None,
) -> list[dict]:
    """Возвращает историю изменений карточки Trello.

    Видно, когда карточку двигали между списками, кто менял срок и требования.

    Аргументы:
        card: URL карточки, shortLink или полный id.
        limit: сколько событий вернуть. По умолчанию 1000.
        types: необязательный фильтр по типам действий через запятую
            (напр. "updateCard,commentCard"). По умолчанию — все типы.

    Возвращает массив объектов: {type, date, who, data}, от новых к старым.
    Полезное в data для перемещений: listBefore / listAfter.
    """
    card = _normalize_card(card)
    params = {"filter": types or "all", "limit": limit}
    actions = await _api_get(f"/cards/{card}/actions", **params)
    return [
        {
            "type": a.get("type"),
            "date": a.get("date"),
            "who": a.get("memberCreator", {}).get("fullName", "—"),
            "data": a.get("data", {}),
        }
        for a in actions
    ]


@mcp.tool()
async def trello_bulk_fetch_board(
    board: str,
    download: bool = False,
    dest_dir: str = "./trello_files",
    include_comments: bool = True,
    max_cards: int = 300,
) -> dict:
    """Собирает по ВСЕЙ доске Trello: карточки, их вложения и комментарии — одним вызовом.

    Для аудита доски. По умолчанию НЕ скачивает файлы (только опись + комментарии),
    т.к. вся доска может весить гигабайты. С download=True скачивает загруженные
    вложения (в подпапки по карточкам, с кэшем); внешние ссылки пропускает.

    Аргументы:
        board: URL доски, shortLink или id.
        download: True — скачать файлы на диск. False — только опись (по умолчанию).
        dest_dir: куда качать. По умолчанию "./trello_files".
        include_comments: включать ли комментарии карточек. По умолчанию True.
        max_cards: ограничение числа карточек в ответе (по умолчанию 300).

    Возвращает: {board, cards, attachments_total, downloaded, items:[{card, url,
    attachments:[{id,name,mime,bytes,is_upload,url,(path,cached)}], comments:[...]}]}.
    Совет: для лёгкой описи ставь download=False; при большом объёме — include_comments=False.
    """
    _require_creds()
    board = _normalize_board(board)

    cards = await _api_get(
        f"/boards/{board}/cards", fields="name,shortLink,url",
        attachments="true", attachment_fields="all", limit=1000,
    )

    comments_by_card: dict[str, list] = {}
    if include_comments:
        actions = await _api_get(
            f"/boards/{board}/actions", filter="commentCard", limit=1000
        )
        for a in actions:
            cid = ((a.get("data") or {}).get("card") or {}).get("id")
            if not cid:
                continue
            comments_by_card.setdefault(cid, []).append({
                "author": a.get("memberCreator", {}).get("fullName", "—"),
                "date": a.get("date"),
                "text": (a.get("data") or {}).get("text", ""),
            })

    items = []
    total_att = downloaded = 0
    jobs = []  # (entry, card_key, meta, dest) для параллельного скачивания
    for c in cards[:max_cards]:
        card_key = c.get("shortLink") or c.get("id")
        atts = []
        for a in c.get("attachments") or []:
            is_up = bool(a.get("isUpload"))
            entry = {
                "id": a.get("id"),
                "name": a.get("fileName") if is_up else a.get("name"),
                "mime": a.get("mimeType"),
                "bytes": a.get("bytes"),
                "is_upload": is_up,
                "url": a.get("url"),
            }
            if download and is_up:
                jobs.append((entry, card_key, a, str(Path(dest_dir) / card_key)))
            atts.append(entry)
        total_att += len(atts)
        items.append({
            "card": c.get("name"),
            "shortLink": c.get("shortLink"),
            "url": c.get("url"),
            "attachments": atts,
            "comments": comments_by_card.get(c.get("id"), []),
        })

    if jobs:
        # Параллельное скачивание с ограничением (лимиты Trello ~300 req/10s).
        sem = asyncio.Semaphore(6)
        async with httpx.AsyncClient(timeout=httpx.Timeout(300.0, connect=30.0)) as client:
            async def _run(entry, ck, meta, dest):
                async with sem:
                    try:
                        p, cached = await _download_attachment(client, ck, meta, dest)
                        entry["path"] = str(p)
                        entry["cached"] = cached
                    except Exception as e:  # noqa: BLE001
                        entry["error"] = str(e)
            await asyncio.gather(*(_run(*j) for j in jobs))
        downloaded = sum(1 for j in jobs if "path" in j[0])

    return {
        "board": board,
        "cards": len(items),
        "cards_total": len(cards),
        "truncated": len(cards) > max_cards or len(cards) >= 1000,
        "attachments_total": total_att,
        "downloaded": downloaded if download else 0,
        "items": items,
    }


@mcp.tool()
async def trello_extract_archive(
    card: str,
    attachment_id: str,
    dest_dir: str = "./trello_files",
) -> dict:
    """Скачивает вложение-архив и распаковывает его. Возвращает список файлов внутри.

    ZIP распаковывается штатно. Для .rar/.7z нужен внешний распаковщик (7-Zip/unrar) —
    в этом случае возвращается путь к скачанному архиву и пометка.

    Аргументы:
        card: URL карточки, shortLink или id.
        attachment_id: id вложения-архива.
        dest_dir: куда скачать/распаковать. По умолчанию "./trello_files".

    Возвращает: {"path", "unpacked_to", "count", "files":[...до 200], "cached"} для zip;
    для не-zip архивов — {"path", "note", "cached"}.
    """
    _require_creds()
    import zipfile
    card = _normalize_card(card)
    async with httpx.AsyncClient(timeout=httpx.Timeout(300.0, connect=30.0)) as client:
        meta = await _get_json(
            client, f"/cards/{card}/attachments/{attachment_id}", fields="all"
        )
        if not meta.get("isUpload", False):
            return {"is_link": True, "url": meta.get("url"), "name": meta.get("name")}
        path, cached = await _download_attachment(client, card, meta, dest_dir)

    if path.suffix.lower() == ".zip":
        target = path.with_name(path.stem + "_unpacked")
        with zipfile.ZipFile(path) as z:
            z.extractall(target)
        files = sorted(
            str(p.relative_to(target)) for p in target.rglob("*")
            if p.is_file() and "__MACOSX" not in str(p)
        )
        return {
            "path": str(path), "unpacked_to": str(target),
            "count": len(files), "files": files[:200], "cached": cached,
        }
    if path.suffix.lower() in (".rar", ".7z"):
        return {
            "path": str(path), "cached": cached,
            "note": f"{path.suffix} требует внешнего распаковщика (7-Zip/unrar). "
                    f"ZIP плагин распаковывает сам.",
        }
    return {"path": str(path), "cached": cached, "note": "не архив"}


@mcp.tool()
async def trello_contact_sheet(
    images_dir: str,
    out_path: str | None = None,
    cols: int = 5,
    cell: int = 320,
    max_images: int = 60,
) -> dict:
    """Собирает контактный лист (montage) из картинок в папке — чтобы обозреть
    большой набор фото одним изображением (дёшево по токенам). Возвращает путь.

    Аргументы:
        images_dir: папка с картинками (напр. распакованный архив съёмки).
        out_path: куда сохранить лист. По умолчанию <images_dir>/_contact_sheet.png.
        cols: колонок в сетке (по умолчанию 5).
        cell: размер ячейки в px (по умолчанию 320).
        max_images: сколько картинок показать (равномерная выборка, по умолчанию 60).

    Возвращает: {"path", "images_total", "shown"}.
    """
    import math
    from PIL import Image
    d = Path(images_dir)
    if not d.exists():
        raise TrelloError(f"Папка не найдена: {images_dir}")
    exts = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif")
    imgs = sorted(
        p for p in d.rglob("*")
        if p.suffix.lower() in exts and "__MACOSX" not in str(p) and "_contact_sheet" not in p.name
    )
    if not imgs:
        raise TrelloError(f"Картинок не найдено в {images_dir}")
    step = max(1, len(imgs) // max_images)
    sample = imgs[::step][:max_images]

    rows = math.ceil(len(sample) / cols)
    pad = 6
    W = cols * (cell + pad) + pad
    H = rows * (cell + pad) + pad
    canvas = Image.new("RGB", (W, H), (238, 238, 240))
    for i, p in enumerate(sample):
        r, c = divmod(i, cols)
        x = pad + c * (cell + pad)
        y = pad + r * (cell + pad)
        try:
            im = Image.open(p).convert("RGB")
        except Exception:
            continue
        im.thumbnail((cell, cell))
        canvas.paste(im, (x + (cell - im.width) // 2, y + (cell - im.height) // 2))
    out = Path(out_path) if out_path else d / "_contact_sheet.png"
    canvas.save(out)
    return {"path": str(out), "images_total": len(imgs), "shown": len(sample)}


def _extract_text(path: Path):
    """Возвращает (текст, вид) из файла или (None, суффикс) для нетекстовых.
    PDF/docx читаются библиотеками — не зависит от того, умеет ли среда их открыть."""
    suf = path.suffix.lower()
    if suf == ".pdf":
        try:
            import pymupdf
        except ImportError as e:
            raise TrelloError(
                "Для чтения PDF нужен pymupdf. Установи зависимости: "
                "pip install -r mcp/requirements.txt"
            ) from e
        doc = pymupdf.open(path)
        return ("\n".join(pg.get_text() for pg in doc), "pdf")
    if suf == ".docx":
        try:
            import docx
        except ImportError as e:
            raise TrelloError(
                "Для чтения docx нужен python-docx. Установи зависимости: "
                "pip install -r mcp/requirements.txt"
            ) from e
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return ("\n".join(parts), "docx")
    if suf in (".txt", ".md", ".csv", ".json", ".log", ".yaml", ".yml"):
        return (path.read_text(encoding="utf-8", errors="replace"), suf.lstrip("."))
    return (None, suf)


@mcp.tool()
async def trello_read_attachment_text(
    card: str,
    attachment_id: str,
    max_chars: int = 20000,
    dest_dir: str = "./trello_files",
) -> dict:
    """Скачивает вложение и возвращает извлечённый ТЕКСТ (для брифов: PDF, docx, txt/md/csv).

    Работает на любой машине, не завися от того, умеет ли среда открывать PDF/docx.
    Для картинок используй trello_fetch_attachment и посмотри файл глазами.

    Аргументы:
        card: URL карточки, shortLink или полный id.
        attachment_id: id вложения (см. trello_list_attachments).
        max_chars: максимум символов в ответе (по умолчанию 20000, чтобы не забить контекст).
        dest_dir: куда сохранить файл. По умолчанию "./trello_files".

    Возвращает для текстового файла: {"path", "name", "kind", "chars", "truncated", "text"}.
    Возвращает для нетекстового:      {"path", "name", "mime", "extractable": false, "note"}.
    Возвращает для внешней ссылки:    {"is_link": true, "url", "name", "text": null}.
    """
    _require_creds()
    card = _normalize_card(card)

    async with httpx.AsyncClient(timeout=httpx.Timeout(300.0, connect=30.0)) as client:
        meta = await _get_json(
            client, f"/cards/{card}/attachments/{attachment_id}", fields="all"
        )
        if not meta.get("isUpload", False):
            return {
                "is_link": True,
                "url": meta.get("url"),
                "name": meta.get("name"),
                "text": None,
            }
        path, cached = await _download_attachment(client, card, meta, dest_dir)
        file_name = path.name

    text, kind = _extract_text(path)
    if text is None:
        return {
            "path": str(path),
            "name": file_name,
            "mime": meta.get("mimeType"),
            "extractable": False,
            "cached": cached,
            "note": (
                f"Тип «{kind}» не текстовый — файл скачан по указанному пути, "
                f"открой его как файл (картинку — посмотри глазами, архив — распакуй)."
            ),
        }
    return {
        "path": str(path),
        "name": file_name,
        "kind": kind,
        "cached": cached,
        "chars": len(text),
        "truncated": len(text) > max_chars,
        "text": text[:max_chars],
    }


# --------------------------------------------------------------------------
# Этап 3: пишущие инструменты (необратимы силами плагина — требуют confirm)
# --------------------------------------------------------------------------
#
# Удаления в плагине нет намеренно: инструментов, стирающих карточки,
# комментарии или вложения, здесь быть не должно. Поэтому всё, что публикуется,
# плагин откатить не может. Оба инструмента защищены параметром confirm:
# без confirm=True они НЕ пишут, а возвращают предпросмотр действия. Агент
# обязан сначала показать этот предпросмотр пользователю и получить явное «да».

COMMENT_MAX = 16384          # лимит Trello на длину комментария
MAX_UPLOAD = 250 * 1024 * 1024  # жёсткий максимум вложения Trello (платный тариф)


@mcp.tool()
async def trello_add_comment(card: str, text: str, confirm: bool = False) -> dict:
    """Публикует комментарий на карточке Trello от имени владельца токена.

    ВНИМАНИЕ: действие необратимо силами плагина (удаления в плагине нет).
    Упоминания через @username порождают реальные уведомления живым людям —
    не рассылай их без явной просьбы пользователя.

    Механизм подтверждения: при confirm=False (по умолчанию) инструмент НИЧЕГО
    не публикует, а возвращает предпросмотр. Покажи предпросмотр пользователю,
    получи явное согласие и только тогда вызови повторно с confirm=True.

    Аргументы:
        card: URL карточки, shortLink или полный id.
        text: текст комментария (поддерживается Markdown). Длиннее 16384
            символов — режется на несколько комментариев, а не теряется.
        confirm: True — публиковать. False — только предпросмотр.

    Возвращает при confirm=False: {"pending_confirmation": true, ...предпросмотр}.
    Возвращает при confirm=True:  {"posted": N, "comments": [{id, date, author}]}.
    """
    _require_creds()
    if not text or not text.strip():
        raise TrelloError("Пустой текст комментария.")
    card_n = _normalize_card(card)
    chunks = [text[i : i + COMMENT_MAX] for i in range(0, len(text), COMMENT_MAX)]

    if not confirm:
        return {
            "pending_confirmation": True,
            "action": "add_comment",
            "card": card_n,
            "parts": len(chunks),
            "chars": len(text),
            "mentions": bool(re.search(r"(?:^|\s)@[\w-]+", text)),
            "preview": text[:500],
            "note": (
                "Комментарий НЕ опубликован. Необратимо силами плагина. "
                "@упоминания шлют реальные уведомления людям. "
                "После явного согласия пользователя вызови снова с confirm=True."
            ),
        }

    created = []
    for chunk in chunks:
        res = await _api_post(
            f"/cards/{card_n}/actions/comments", data={"text": chunk}
        )
        created.append(
            {
                "id": res.get("id"),
                "date": res.get("date"),
                "author": res.get("memberCreator", {}).get("fullName"),
            }
        )
    return {"posted": len(created), "comments": created}


@mcp.tool()
async def trello_upload_attachment(
    card: str,
    file_path: str,
    name: str | None = None,
    confirm: bool = False,
) -> dict:
    """Прикрепляет локальный файл к карточке Trello.

    ВНИМАНИЕ: действие необратимо силами плагина (удаления в плагине нет).

    Механизм подтверждения: при confirm=False (по умолчанию) инструмент НИЧЕГО
    не прикрепляет, а возвращает предпросмотр. Покажи его пользователю, получи
    явное согласие и только тогда вызови повторно с confirm=True.

    Аргументы:
        card: URL карточки, shortLink или полный id.
        file_path: путь к локальному файлу (проверяется до отправки).
        name: имя вложения. По умолчанию — имя файла.
        confirm: True — прикреплять. False — только предпросмотр.

    Лимит размера зависит от тарифа рабочего пространства: 10 МБ на бесплатном,
    до 250 МБ на платных. Файл больше 250 МБ отклоняется до отправки.

    Возвращает при confirm=False: {"pending_confirmation": true, ...предпросмотр}.
    Возвращает при confirm=True:  {id, name, bytes, mime, url}.
    """
    _require_creds()
    p = Path(file_path)
    if not p.exists() or not p.is_file():
        raise TrelloError(f"Файл не найден: {file_path}")
    size = p.stat().st_size
    if size > MAX_UPLOAD:
        raise TrelloError(
            f"Файл {size / 1e6:.1f} МБ превышает максимум Trello 250 МБ."
        )
    card_n = _normalize_card(card)
    att_name = name or p.name

    if not confirm:
        return {
            "pending_confirmation": True,
            "action": "upload_attachment",
            "card": card_n,
            "file": str(p),
            "name": att_name,
            "bytes": size,
            "note": (
                f"Файл НЕ прикреплён. Необратимо силами плагина. "
                f"Размер {size / 1e6:.1f} МБ (на бесплатном тарифе лимит 10 МБ, "
                f"на платных до 250 МБ). После согласия пользователя вызови "
                f"снова с confirm=True."
            ),
        }

    mime = mimetypes.guess_type(str(p))[0] or "application/octet-stream"
    with open(p, "rb") as fh:
        res = await _api_post(
            f"/cards/{card_n}/attachments",
            data={"name": att_name},
            files={"file": (att_name, fh, mime)},
        )
    return {
        "id": res.get("id"),
        "name": res.get("name"),
        "bytes": res.get("bytes"),
        "mime": res.get("mimeType"),
        "url": res.get("url"),
    }


if __name__ == "__main__":
    mcp.run()
